"""
文档输出模块 v2 —— 生成标准格式的投标文件 DOCX
按照政府采购文件标准格式排版
"""
import os
import re
from typing import Dict

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


class BidDocumentWriter:

    def __init__(self):
        if not _DOCX_OK:
            raise ImportError("请安装: pip install python-docx")
        self.doc = Document()
        self._init_styles()

    def _init_styles(self):
        """初始化文档样式，符合正式投标文件标准"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '仿宋'
        font.size = Pt(14)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        pf = style.paragraph_format
        pf.line_spacing = 1.5
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

        for sec in self.doc.sections:
            sec.top_margin = Cm(2.54)
            sec.bottom_margin = Cm(2.54)
            sec.left_margin = Cm(3.17)
            sec.right_margin = Cm(3.17)

    # ===== 基础操作 =====

    def _run(self, text: str, font_name='仿宋', font_size=14, bold=False, alignment=None, color=None):
        """添加段落"""
        p = self.doc.add_paragraph()
        if alignment is not None:
            p.alignment = alignment
        run = p.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        return p

    def _title(self, text: str, level=0):
        """标题：黑体加粗"""
        sizes = {0: 22, 1: 16, 2: 15}
        sz = sizes.get(level, 14)
        p = self._run(text, font_name='黑体', font_size=sz, bold=True)
        if level == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def _empty_line(self, n=1):
        for _ in range(n):
            self.doc.add_paragraph()

    def _page_break(self):
        self.doc.add_page_break()

    def _table(self, headers: list, rows: list, col_widths=None):
        """添加表格"""
        tbl = self.doc.add_table(rows=len(rows) + 1, cols=len(headers), style='Table Grid')
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(12)
                    r.font.name = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = tbl.rows[i + 1].cells[j]
                cell.text = str(val)
                for r in cell.paragraphs[0].runs:
                    r.font.size = Pt(12)
                    r.font.name = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        self._empty_line()
        return tbl

    # ===== 核心构建 =====

    def build(self, bid_document: Dict, project_info: Dict) -> str:
        """
        bid_document: { chapter_key: { 'title': ..., 'content': ... }, ... }
        """
        first_key = None
        cover_done = False

        for chapter_key, chapter_data in bid_document.items():
            if first_key is None:
                first_key = chapter_key

            title = chapter_data.get('title', '')
            content = chapter_data.get('content', '')

            # 封面特殊处理
            if chapter_key == 'cover':
                self._render_cover(content, project_info)
                cover_done = True
                self._page_break()
                continue

            # 其他章节
            self._render_chapter(title, content)
            self._page_break()

        # 如果缺少封面，自动生成
        if not cover_done:
            self._render_cover_auto(project_info)
            self._page_break()
            # 把自动封面插入最前面比较困难，这里直接附加在开头需要重建
            # 简单方案：用 cover 的标题重新构建
            pass

        # 保存
        output_dir = os.path.join(
            os.path.dirname(os.path.dirname(__file__)), 'output'
        )
        os.makedirs(output_dir, exist_ok=True)

        proj_name = project_info.get('project_name', 'bid')
        proj_id = project_info.get('project_id', '')
        safe_name = re.sub(r'[\\/:*?"<>|]', '_', f"{proj_name}_{proj_id}")
        output_path = os.path.join(output_dir, f'{safe_name}_投标文件.docx')
        self.doc.save(output_path)
        return output_path

    def _render_cover(self, text: str, project_info: Dict):
        """渲染封面 —— 解析生成器输出的封面内容"""
        lines = text.split('\n')

        for line in lines:
            stripped = line.strip()
            if not stripped:
                self._empty_line()
                continue

            # 分隔线
            if set(stripped) == {'='} or stripped.startswith('==='):
                self._run('─' * 40, font_name='宋体', font_size=12,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
                continue

            # 正本/副本
            if '正本' in stripped or '副本' in stripped:
                self._run(stripped, font_name='宋体', font_size=16,
                          alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                continue

            # 项目名称行（居中大字）
            if ('项目' in stripped or '工程' in stripped or '采购' in stripped or '服务' in stripped) \
               and '编号' not in stripped:
                if re.search(r'(?:项目|工程|采购|服务)', stripped):
                    self._title(stripped, level=0)
                    continue

            # "投标文件"行
            if '投标' in stripped and '文件' in stripped and '人' not in stripped:
                self._title(stripped, level=0)
                self._empty_line(2)
                continue

            # 项目编号
            if '编号' in stripped:
                self._run(stripped, font_name='仿宋', font_size=16,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
                continue

            # 签署行
            if any(kw in stripped for kw in ['投标人', '法定代表人', '授权代表', '日期', '招标人']):
                self._run(stripped, font_name='仿宋', font_size=16, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
                continue

            # 封面标注信息
            if stripped.startswith('·') or stripped.startswith('-'):
                self._run(stripped, font_name='仿宋', font_size=12,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
                continue

            # 其他
            self._run(stripped, font_name='仿宋', font_size=14,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def _render_cover_auto(self, project_info: Dict):
        """兜底自动封面"""
        proj_name = project_info.get('project_name', '________项目')
        proj_id = project_info.get('project_id', '________')
        bidder = project_info.get('bidder_name', '')
        bid_date = project_info.get('bid_opening_date', project_info.get('submission_date', ''))
        tenderer = project_info.get('tenderee_name', '')

        self._run('正本／副本', font_size=16, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        self._empty_line(4)
        self._title(proj_name, level=0)
        if proj_id and proj_id != '________':
            self._run(f'（项目编号：{proj_id}）', font_size=16,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._empty_line(2)
        self._title('投  标  文  件', level=0)
        self._empty_line(6)
        if tenderer and tenderer != '贵单位':
            self._run(f'招标人：{tenderer}', font_size=16, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._run(f'投标人：{bidder}（公章）', font_size=16, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._run('法定代表人或授权代表：______________（签字）', font_size=16, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if bid_date:
            self._run(f'日  期：{bid_date}', font_size=16, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def _render_chapter(self, title: str, content: str):
        """渲染一个章节"""
        # 章节标题：黑体加粗
        clean_title = re.sub(r'^[\d]+\s*[\.、）\)]\s*', '', title).strip()
        if not clean_title:
            clean_title = title.strip()
        self._title(clean_title, level=1)
        self._empty_line()

        # 内容
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                self._empty_line()
                continue

            # 子标题
            if re.match(r'^[一二三四五六七八九十]、', stripped):
                self._title(stripped, level=2)
                continue
            if re.match(r'^[（\(][一二三四五六七八九十]+[）\)]', stripped):
                self._title(stripped, level=2)
                continue

            # 表格行
            if stripped.startswith('|') and stripped.endswith('|'):
                continue  # 复杂表格暂时跳过，保持纯文本

            # 分隔线
            if set(stripped) in ({'─'}, {'='}, {'━'}):
                continue

            # 普通段落
            self._run(stripped)


def generate_docx(bid_document: Dict, project_info: Dict, output_path: str = None) -> str:
    writer = BidDocumentWriter()
    path = writer.build(bid_document, project_info)
    if output_path:
        import shutil
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        shutil.copy(path, output_path)
        return output_path
    return path


if __name__ == "__main__":
    print("文档输出模块已就绪，请通过 process_all.py 或 main.py 调用。")
