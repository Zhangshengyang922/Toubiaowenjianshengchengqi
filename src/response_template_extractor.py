"""
响应文件模板提取器 v4 —— 直接从招标文件DOCX中提取"响应文件格式"章节
使用XML级别替换占位符，完全保留原始格式不变

v4 改进：消除所有硬编码位置假设，适应任意招标文件结构
- 章节定位：全文搜索，取最后匹配（避免目录干扰），不做段落索引跳过
- 封面检测：自动感知格式分界（第二个独立格式标记 = 封面结束）
- 目录生成：通用正则匹配格式/附件/表格等任意格式标识
- 目录分组：自动按前缀分组，不硬编码"格式1-/格式2-"
"""
import os
import re
import copy
from lxml import etree
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Word文档XML命名空间
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

# ── 通用正则：匹配章节标题 + 格式编号，不硬编码章节号/名称 ──
# 要求段落以"第X章"开头（真正的章节标题），限制间距 .{0,30}? 防止误匹配
_CHAPTER_PATTERN = re.compile(
    r'^\s*第[一二三四五六七八九十\d]+章.{0,30}?(?:响应文件格式|投标文件格式|格式要求|响应文件)'
)

# 格式编号：格式X-X / 附件X / 附表X / 表格X / 附录X（支持全角半角连字符）
_FORMAT_MARKER_PATTERN = re.compile(
    r'(?:格式|附件|附表|表格|附录)\s*[零一二三四五六七八九十\d]+[-–—.\u2013\u2014\u2015]?[零一二三四五六七八九十\d]*'
)

# ── 用于封面截止检测：按顺序记录所有出现的独立格式标记 ──
_UNIQUE_FORMAT_RE = re.compile(
    r'((?:格式|附件|附表|表格|附录)\s*[零一二三四五六七八九十\d]+[-–—.\u2013\u2014\u2015]?[零一二三四五六七八九十\d]*)'
)

# ── 标题自动识别正则 ──
# 一级标题：格式X-Y、附件X、附表X 等（响应文件的主要章节）
_LEVEL1_FORMAT_RE = re.compile(
    r'^\s*(?:格式|附件|附表|表格|附录)\s*[零一二三四五六七八九十\d]+[-–—·.\u2013\u2014\u2015]?[零一二三四五六七八九十\d]*'
)
# 一级标题：一、 二、 三、 等中文序号开头（加粗且较短，用于格式章节内的一级标题）
_LEVEL1_CN_RE = re.compile(r'^\s*[一二三四五六七八九十]+[、，．.]\s*\S')
# 二级标题：（一）（二）（1）（2）等括号序号开头
_LEVEL2_PAREN_RE = re.compile(r'^\s*[（(]\s*[一二三四五六七八九十\d]+\s*[）)]\s*\S')
# 排除：注、说明、注意等不是标题
_NOT_HEADING_RE = re.compile(
    r'^\s*(?:注\s*[：:意]|说明\s*[：:]|注意\s*[：:事项]|特别提醒|重要提示|提示\s*[：:]|备注\s*[：:]|※|致\s*[：:])'
)
# 排除：项目编号、供应商盖章、日期、签名、附表、纯序号等
_EXCLUDE_PATTERNS = [
    re.compile(r'^\s*(?:项目编号|项目名称)\s*[：:]'),
    re.compile(r'^\s*供应商\s*[：:].*(?:盖章|公章|签字)'),
    re.compile(r'^\s*日\s*期\s*[：:]'),
    re.compile(r'^\s*附表\s*[：:]?\s*$'),
    re.compile(r'^\s*最后报价表\s*$'),
    re.compile(r'^\s*[一二三四五六七八九十\d]+\s*[、，,\.．]\s*$'),
]


class ResponseTemplateExtractor:
    """从招标文件DOCX提取响应文件格式章节，填充后生成响应文件"""

    def __init__(self):
        self.paragraph_start = None
        self.paragraph_end = None
        self.chapter_keyword = None  # 实际匹配到的章节关键词（如"第七章 响应文件格式"）

    def generate(self, docx_path: str, fill_data: dict) -> str:
        """
        docx_path: 招标文件DOCX路径
        fill_data: 填充数据字典
        """
        doc = Document(docx_path)
        self._find_format_chapter(doc)
        if self.paragraph_start is None:
            raise ValueError(
                "未找到包含'响应文件格式'的章节。"
                "请确认招标文件中存在类似'第X章 响应文件格式'的标题。"
            )

        new_doc = self._build_new_document(doc, fill_data)

        output_dir = os.path.join(
            os.path.dirname(os.path.dirname(__file__)), 'output'
        )
        os.makedirs(output_dir, exist_ok=True)
        safe_name = re.sub(r'[\\/:*?"<>|]', '_', fill_data.get('project_name', 'response'))
        output_path = os.path.join(output_dir, f'{safe_name}_响应文件.docx')
        new_doc.save(output_path)
        return output_path

    # ==================== 定位格式章节（无硬编码位置假设） ====================

    def _find_format_chapter(self, doc):
        """
        全文搜索找到"响应文件格式"对应的章节范围。

        策略（无段落索引硬编码）：
        1. 扫描所有段落，匹配 "第X章" + "格式/响应文件" 模式（支持任意章节号）
        2. 取最后匹配的作为真实章节标题（目录中的引用总是出现在前面）
        3. 自动查找下一章作为结束边界
        """
        all_matches = []  # [(paragraph_index, chapter_number_text, full_text)]

        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            # 匹配 "第X章" + 格式相关描述
            m = _CHAPTER_PATTERN.search(text)
            if m:
                # 提取章节数字用于寻找下一章
                chap_num_match = re.search(r'第([一二三四五六七八九十\d]+)章', text)
                chap_num = chap_num_match.group(1) if chap_num_match else ''
                all_matches.append((i, chap_num, text))

        if not all_matches:
            return

        # ── 取最后匹配的作为真实章节标题 ──
        # 原因：招标文件的目录中会先出现章节引用，真实章节内容在后面
        #      取最后一个匹配是最安全的策略
        last_idx, chap_num, full_text = all_matches[-1]
        self.paragraph_start = last_idx
        self.chapter_keyword = full_text[:60]

        # ── 查找下一章作为结束标记 ──
        next_chap_num = self._next_chapter_number(chap_num)
        if next_chap_num:
            for i, p in enumerate(doc.paragraphs):
                if i <= self.paragraph_start:
                    continue
                text = p.text.strip()
                if f'第{next_chap_num}章' in text:
                    self.paragraph_end = i
                    break

        # 如果找不到下一章，取文档末尾
        if self.paragraph_end is None:
            self.paragraph_end = len(doc.paragraphs)

        print(f"    [定位] 章节: {self.chapter_keyword} (段落{self.paragraph_start}~{self.paragraph_end})")

    @staticmethod
    def _next_chapter_number(chap_num: str) -> str | None:
        """计算下一章的数字（支持中文数字和阿拉伯数字）"""
        cn_map = {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':10,
                  '十一':11,'十二':12,'十三':13,'十四':14,'十五':15}
        num = None
        if chap_num.isdigit():
            num = int(chap_num)
        elif chap_num in cn_map:
            num = cn_map[chap_num]
        if num is not None:
            next_num = num + 1
            # 优先用中文数字
            rev_map = {v:k for k,v in cn_map.items()}
            return rev_map.get(next_num, str(next_num))
        return None

    # ==================== 构建新文档 ====================

    def _build_new_document(self, src_doc, fill_data):
        """按 body 元素顺序遍历，同时处理段落和表格"""
        new_doc = Document()

        for sec in new_doc.sections:
            sec.top_margin = Cm(2.5)
            sec.bottom_margin = Cm(2.5)
            sec.left_margin = Cm(2.5)
            sec.right_margin = Cm(2.5)


        body = new_doc.element.body
        src_body = src_doc.element.body

        # 第一步：按 body 元素顺序收集格式章节的所有元素
        para_count = 0
        table_count = 0
        ch7_elements = []

        for child in src_body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                if para_count < self.paragraph_start:
                    para_count += 1
                    continue
                if para_count >= self.paragraph_end:
                    break

                cloned = copy.deepcopy(child)
                self._replace_in_xml(cloned, fill_data)
                ch7_elements.append(cloned)
                para_count += 1

            elif tag == 'tbl':
                # 表格位于段落之间，判断是否在章节范围内
                if self.paragraph_start <= para_count < self.paragraph_end:
                    cloned = copy.deepcopy(child)
                    self._replace_in_xml(cloned, fill_data)
                    ch7_elements.append(cloned)
                    table_count += 1

        # ── 封面检测：定位封面起止 ──
        cover_start, cover_end = self._detect_cover_end(ch7_elements)
        if cover_start < 0:
            cover_start = 0
            cover_end = 0

        # ── 清除所有元素的旧 outlineLvl（源文档可能残留）──
        for elem in ch7_elements:
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            if tag == 'p':
                pPr = elem.find(f'{{{W}}}pPr')
                if pPr is not None:
                    for old in pPr.findall(f'{{{W}}}outlineLvl'):
                        pPr.remove(old)

        print(
            f"    [调试] 共 {len(ch7_elements)} 个元素（含 {table_count} 个表格），"
            f"封面索引: {cover_start}~{cover_end}"
        )

        # 第二步：封面 → 分页 → 目录 → 分页 → 正文
        # 2.1 跳过"第X章 响应文件格式"及编制须知，只保留封面
        #    封面内除签名区（供应商/法定代表人/日期/签字/盖章/授权）外，统一按标题格式设置
        #    封面标题（项目名称行）设为一级标题 outlineLvl=0
        project_name = fill_data.get('project_name', '')
        signature_keywords = ['供应商', '法定代表人', '签字', '盖章', '授权']
        for i in range(cover_start, cover_end):
            elem = ch7_elements[i]
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            if tag == 'p':
                text = self._get_element_text(elem)
                is_signature = any(kw in text for kw in signature_keywords) or re.search(r'^\s*日\s*期', text)
                if not is_signature:
                    self._apply_title_format(elem)
                # 封面标题行（含项目名的段落）→ 文档一级标题
                if project_name and project_name in text and len(text.strip()) <= 80:
                    self._set_outline_level(elem, 0)
            body.append(elem)



        # 2.2 分页符 + 目录（封面存在且正文存在时才插入）
        if cover_end > cover_start and cover_end < len(ch7_elements):
            self._add_section_break(body)
            self._add_toc(new_doc, fill_data, ch7_elements, cover_end)

        # 2.3 分页符 + 正文（先识别标题层级，再写入）
        self._classify_headings(ch7_elements, cover_end)
        self._add_section_break(body)
        for i in range(cover_end, len(ch7_elements)):
            body.append(ch7_elements[i])

        return new_doc


    def _detect_cover_end(self, ch7_elements: list) -> tuple:
        """
        自动检测封面起止位置。

        策略：
        1. 先找 "封面" / "封面格式" 关键字，作为封面起点
        2. 再找封面之后的第一个实质性格式/内容标记，作为封面终点
        3. 支持 "格式自拟" 这种无编号但标志内容开始的文本

        返回 (cover_start, cover_end)：
        - cover_start: 封面起始索引（-1 表示未找到封面）
        - cover_end: 封面结束索引（第一个非封面元素）
        """
        cover_start = -1
        for i, elem in enumerate(ch7_elements):
            text = self._get_element_text(elem)
            if re.search(r'封面格式?\s*[:：]?', text):
                cover_start = i
                break

        if cover_start < 0:
            return -1, 0

        # 封面后的第一个实质格式标记
        for i in range(cover_start + 1, len(ch7_elements)):
            text = self._get_element_text(ch7_elements[i])
            if (
                _UNIQUE_FORMAT_RE.search(text)
                or '格式自拟' in text
                or re.search(r'^\s*(?:[一二三四五六七八九十]+[、.．])', text.strip())
            ):
                return cover_start, i

        return cover_start, len(ch7_elements)

    def _apply_title_format(self, p_elem):
        """把段落统一设为宋体、四号、1.5 倍行距"""
        # 设置所有 run 的字体和字号
        for r in p_elem.findall(f'.//{{{W}}}r'):
            rPr = r.find(f'{{{W}}}rPr')
            if rPr is None:
                rPr = etree.Element(f'{{{W}}}rPr')
                r.insert(0, rPr)

            rFonts = rPr.find(f'{{{W}}}rFonts')
            if rFonts is None:
                rFonts = etree.Element(f'{{{W}}}rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), '宋体')
            rFonts.set(qn('w:ascii'), '宋体')

            sz = rPr.find(f'{{{W}}}sz')
            if sz is None:
                sz = etree.Element(f'{{{W}}}sz')
                rPr.append(sz)
            sz.set(qn('w:val'), '28')  # 四号 = 14pt = 28 半磅

            szCs = rPr.find(f'{{{W}}}szCs')
            if szCs is None:
                szCs = etree.Element(f'{{{W}}}szCs')
                rPr.append(szCs)
            szCs.set(qn('w:val'), '28')

        # 设置段落行距 1.5 倍
        pPr = p_elem.find(f'{{{W}}}pPr')
        if pPr is None:
            pPr = etree.Element(f'{{{W}}}pPr')
            p_elem.insert(0, pPr)
        spacing = pPr.find(f'{{{W}}}spacing')
        if spacing is None:
            spacing = etree.Element(f'{{{W}}}spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), '360')      # 1.5 倍行距
        spacing.set(qn('w:lineRule'), 'auto')


    def _get_element_text(self, element):
        """获取XML元素中所有文本"""
        texts = []
        for t in element.findall(f'.//{{{W}}}t'):
            if t.text:
                texts.append(t.text)
        return ''.join(texts)

    # ==================== 标题自动识别：一级/二级标题 ====================

    def _classify_headings(self, elements: list, start_idx: int = 0):
        """
        自动识别正文中的一级标题和二级标题，并添加 outlineLvl 属性。

        策略（纯文本+格式特征）：

        一级标题 (outlineLvl=0)：
        - "格式X-Y / 附件X / 附表X" 等 —— 响应文件的主格式标记
        - "一、xxx / 二、xxx" —— 中文序号开头、加粗、且简短（≤25字）

        二级标题 (outlineLvl=1)：
        - "（一）/（二）/（1）/（2）" 等括号序号开头、加粗、且简短（≤20字）

        排除规则：
        - "注/说明/注意/重要提示/致/备注" 开头的不算标题
        - 项目编号行、供应商盖章行、日期行、附表行不算标题
        - 纯序号行（如 "1、" 后面无实质内容）不算标题
        """
        # 重新识别标题（旧 outlineLvl 已在 _build_new_document 中统一清除）
        for i in range(start_idx, len(elements)):
            elem = elements[i]
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            if tag != 'p':
                continue

            text = self._get_element_text(elem).strip()
            if not text:
                continue

            # 排除非标题模式
            if _NOT_HEADING_RE.match(text):
                continue
            if any(pat.search(text) for pat in _EXCLUDE_PATTERNS):
                continue

            is_bold = self._element_has_bold(elem)
            level = None

            # ── 规则1：格式/附件/附表 → 一级标题 ──
            if _LEVEL1_FORMAT_RE.match(text):
                level = 0

            # ── 规则2：一、二、三、... 中文序号 → 一级标题（加粗 + ≤25字） ──
            elif _LEVEL1_CN_RE.match(text) and is_bold and len(text) <= 25:
                level = 0

            # ── 规则3：（一）（二）（1）（2）... 括号序号 → 二级标题（加粗 + ≤20字） ──
            elif _LEVEL2_PAREN_RE.match(text) and is_bold and len(text) <= 20:
                level = 1

            if level is not None:
                self._set_outline_level(elem, level)
                if not is_bold:
                    self._ensure_bold(elem)

    @staticmethod
    def _element_has_bold(element) -> bool:
        """检查元素中是否有加粗段"""
        for r in element.findall(f'.//{{{W}}}r'):
            rPr = r.find(f'{{{W}}}rPr')
            if rPr is not None and rPr.find(f'{{{W}}}b') is not None:
                return True
        return False

    @staticmethod
    def _element_is_all_bold(element) -> bool:
        """检查元素是否所有文本段都加粗（排除空段）"""
        has_text = False
        for r in element.findall(f'.//{{{W}}}r'):
            t = r.find(f'{{{W}}}t')
            if t is not None and (t.text or '').strip():
                has_text = True
                rPr = r.find(f'{{{W}}}rPr')
                if rPr is None or rPr.find(f'{{{W}}}b') is None:
                    return False
        return has_text

    @staticmethod
    def _is_signature_or_date(text: str) -> bool:
        """判断文本是否为签名/日期行"""
        return bool(
            re.search(r'(?:供应商|法定代表人|授权代表|签字|盖章|日期[：:]|时间[：:]|年\s*月\s*日)', text)
        )

    @staticmethod
    def _set_outline_level(element, level: int):
        """给段落元素设置 outlineLvl"""
        pPr = element.find(f'{{{W}}}pPr')
        if pPr is None:
            pPr = etree.Element(f'{{{W}}}pPr')
            element.insert(0, pPr)
        # 移除已有 outlineLvl
        for old in pPr.findall(f'{{{W}}}outlineLvl'):
            pPr.remove(old)
        ol = etree.Element(f'{{{W}}}outlineLvl')
        ol.set(qn('w:val'), str(level))
        pPr.append(ol)

    @staticmethod
    def _ensure_bold(element):
        """确保元素中所有文本段加粗"""
        for r in element.findall(f'.//{{{W}}}r'):
            rPr = r.find(f'{{{W}}}rPr')
            if rPr is None:
                rPr = etree.Element(f'{{{W}}}rPr')
                r.insert(0, rPr)
            if rPr.find(f'{{{W}}}b') is None:
                b = etree.Element(f'{{{W}}}b')
                rPr.append(b)

    # ==================== 目录生成（无硬编码格式假设） ====================

    def _add_toc(self, new_doc, fill_data, ch7_elements, cover_end):
        """在封面后生成目录页（宋体、四号、1.5 倍行距）"""
        body = new_doc.element.body

        # 收集格式标题及描述（从正文开始，避免封面干扰）
        format_titles = []
        for i, elem in enumerate(ch7_elements[cover_end:], start=cover_end):
            text = self._get_element_text(elem)
            # 匹配 "格式X-X" / "附件X" / "附表X" 等任意格式标记
            m = _FORMAT_MARKER_PATTERN.search(text)
            if m:
                title_key = m.group()
                if title_key not in [t[0] for t in format_titles]:
                    rest = text[m.end():].strip()
                    rest = re.sub(r'^[：:、，,\s]+', '', rest)
                    if not rest and i + 1 < len(ch7_elements):
                        next_text = self._get_element_text(ch7_elements[i + 1]).strip()
                        if next_text and not any(w in next_text for w in ['正本', '副本', '格式', '附件', '附表', '表格']):
                            rest = next_text
                    full_title = f'{title_key}  {rest}' if rest else title_key
                    format_titles.append((title_key, full_title))

        if not format_titles:
            return

        # 去重保持顺序
        seen = set()
        unique_titles = []
        for key, title in format_titles:
            if key not in seen:
                seen.add(key)
                unique_titles.append((key, title))

        # 写入目录标题
        toc_heading = self._make_paragraph(
            '目    录',
            font_name='宋体', font_size=Pt(14), bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5
        )
        body.append(toc_heading)
        body.append(self._make_paragraph('', line_spacing=1.5))

        # ── 自动按前缀分组 ──
        # 提取每个标题的前缀（如 "格式1"、"格式2"、"附件一" 等）
        groups = self._group_titles(unique_titles)

        if len(groups) > 1:
            # 多组 → 分组显示
            group_names = {
                0: '第一部分',
                1: '第二部分',
                2: '第三部分',
                3: '第四部分',
                4: '第五部分',
            }
            for gi, group in enumerate(groups):
                group_label = group_names.get(gi, f'第{gi+1}部分')
                body.append(self._make_paragraph(
                    f'{group_label}',
                    font_name='宋体', font_size=Pt(14), bold=True,
                    left_indent=Cm(0.5), line_spacing=1.5
                ))
                for key, title in group:
                    body.append(self._make_paragraph(
                        title,
                        font_name='宋体', font_size=Pt(14),
                        left_indent=Cm(1.5), line_spacing=1.5
                    ))
                body.append(self._make_paragraph('', line_spacing=1.5))
        else:
            # 单组 → 直接列表
            for key, title in unique_titles:
                body.append(self._make_paragraph(
                    title,
                    font_name='宋体', font_size=Pt(14),
                    left_indent=Cm(1.5), line_spacing=1.5
                ))


    @staticmethod
    def _group_titles(titles: list) -> list:
        """
        按前缀自动分组。
        返回 [[(key, title), ...], ...]
        """
        if not titles:
            return [titles]

        prefix_re = re.compile(r'^((?:格式|附件|附表|表格|附录)\s*\d+)')
        groups = []
        current_prefix = None
        current_group = []

        for key, title in titles:
            pm = prefix_re.match(key)
            prefix = pm.group(1) if pm else key
            if current_prefix is not None and prefix != current_prefix:
                groups.append(current_group)
                current_group = []
            current_prefix = prefix
            current_group.append((key, title))

        if current_group:
            groups.append(current_group)

        return groups if groups else [titles]

    def _make_paragraph(self, text, font_name='仿宋', font_size=None, bold=False,
                        alignment=None, left_indent=None, line_spacing=None):
        """创建格式化的段落"""
        from docx.oxml import OxmlElement

        p = OxmlElement('w:p')

        # 段落属性
        pPr = OxmlElement('w:pPr')
        if alignment is not None:
            jc = OxmlElement('w:jc')
            align_map = {
                WD_ALIGN_PARAGRAPH.CENTER: 'center',
                WD_ALIGN_PARAGRAPH.LEFT: 'left',
                WD_ALIGN_PARAGRAPH.RIGHT: 'right',
            }
            jc.set(qn('w:val'), align_map.get(alignment, 'left'))
            pPr.append(jc)
        if left_indent:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(int(left_indent / Cm(1) * 567)))  # EMU
            pPr.append(ind)
        if line_spacing is not None:
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), str(int(line_spacing * 240)))
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)
        p.append(pPr)

        # 文本run
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if font_name:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:ascii'), font_name)
            rPr.append(rFonts)
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))  # half-points
            rPr.append(sz)
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        r.append(rPr)

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)

        return p


    def _add_section_break(self, body):
        """添加分页符"""
        from docx.oxml import OxmlElement
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        p.append(r)
        body.append(p)

    # ==================== XML级别占位符替换 ====================

    def _replace_in_xml(self, element, fill_data: dict):
        """在XML元素中递归替换所有文本节点中的占位符"""
        text_elements = element.findall(f'.//{{{W}}}t')
        full_text = ''
        text_elems = []
        for t_elem in text_elements:
            text = t_elem.text or ''
            full_text += text
            text_elems.append((t_elem, text))

        if not full_text.strip():
            return

        new_full = self._apply_replacements(full_text, fill_data)
        if new_full == full_text:
            return

        for idx, (t_elem, _) in enumerate(text_elems):
            if idx == 0:
                t_elem.text = new_full
            else:
                t_elem.text = ''

    def _apply_replacements(self, text: str, fill: dict) -> str:
        """对所有占位符进行替换"""
        bidder = fill.get('bidder_name', '') or ''
        legal = fill.get('legal_representative', '') or ''
        authorized = fill.get('authorized_person', '') or ''
        agency = fill.get('agency_name', '') or ''
        proj_name = fill.get('project_name', '') or ''
        proj_id = fill.get('project_id', '') or ''
        date_val = fill.get('bid_opening_date', '') or ''
        address = fill.get('address', '') or ''
        phone = fill.get('phone', '') or ''
        zip_code = fill.get('zip_code', '') or ''
        fax = fill.get('fax', '') or ''

        t = text

        # ==== 项目名称: xxxx项目 / XX项目 ====
        t = re.sub(r'[xX]{2,4}\s*项目', proj_name, t)
        t = re.sub(r'[xX]{2}\s*项目', proj_name, t)
        # 空格型项目名占位：项目名称：                           
        t = re.sub(r'项目名称[：:]\s*[xX _]*\s*$', f'项目名称：{proj_name}', t)
        # （项目名称）
        t = re.sub(r'（\s*[xX _]*项目名称[xX _]*\s*）', f'（{proj_name}）', t)
        # （项目名称：           、项目编号：              ）
        t = re.sub(
            r'项目名称[：:]\s*[xX _]*\s*[、,，]\s*项目编号[：:]\s*[xX _]*\s*[）)]',
            f'项目名称：{proj_name}、项目编号：{proj_id}）',
            t
        )


        # ==== 引号内项目名称 ====
        t = re.sub(r'"X{3,10}"', f'"{proj_name}"', t)
        t = re.sub(r"'X{3,10}'", f"'{proj_name}'", t)
        t = re.sub(r'"[xX]{3,10}"', f'"{proj_name}"', t)
        t = re.sub(r'\u201c[xX]{3,10}\u201d', f'\u201c{proj_name}\u201d', t)

        # ==== 项目编号 ====
        if proj_id:
            t = t.replace('（项目编号：XXXX）', f'（项目编号：{proj_id}）')
            t = t.replace('(项目编号：XXXX)', f'(项目编号：{proj_id})')
            t = t.replace('项目编号：XXXX', f'项目编号：{proj_id}')
            t = re.sub(r'项目编号[：:]\s*[xX_]*\s*$', f'项目编号：{proj_id}', t)
            t = re.sub(r'项目编号[：:]\s*[xX_]*\s*[,，、;；\.。]?\s*$', f'项目编号：{proj_id}', t)
            t = re.sub(r'项目编号[：:]\s*[xX _]*\s*[）)]', f'项目编号：{proj_id}）', t)


        # ==== 日期 ====

        if date_val:
            t = re.sub(r'XXX+年XXX+月XXX+日', date_val, t)
            t = re.sub(r'XX年XX月XX日', date_val, t)
            t = re.sub(r'日\s*期[：:]\s*XXX+年XXX+月XXX+日', f'日    期：{date_val}', t)
            t = re.sub(r'时间[：:]\s*XX年XX月XX日', f'时间：{date_val}', t)
            t = re.sub(r'日\s*期[：:]\s*[xX]{3,4}\s*[。]?', f'日    期：{date_val}', t)
            t = re.sub(r'日期[：:]\s*[xX]{3,4}\s*[。]?', f'日期：{date_val}', t)
            t = re.sub(r'日期[：:]\s*[xX]{3,4}\s*$', f'日期：{date_val}', t)
            # 空格/空值型日期占位：日期：       年     月    日
            t = re.sub(r'日\s*期[：:]\s*[xX _]*\s*年\s*[xX _]*\s*月\s*[xX _]*\s*日', f'日    期：{date_val}', t)
            t = re.sub(r'日期[：:]\s*[xX _]*\s*年\s*[xX _]*\s*月\s*[xX _]*\s*日', f'日期：{date_val}', t)
            t = re.sub(r'时间[：:]\s*[xX _]*\s*年\s*[xX _]*\s*月\s*[xX _]*\s*日', f'时间：{date_val}', t)

            t = re.sub(r'时间[：:]\s*[年xX _]*\s*年\s*[月xX _]*\s*月\s*[日xX _]*\s*日', f'时间：{date_val}', t)


        # ==== 采购代理机构 ====
        if agency:
            t = re.sub(r'[xX]{3,8}\s*[（\(]采购代理机构名称[）\)]\s*[：:]', f'{agency}：', t)
            t = re.sub(r'[xX]{3,4}\s*[（\(]采购代理机构名称[）\)]\s*[：:]', f'{agency}：', t)

        # ==== 供应商名称：XXXX/XXX ====
        if bidder:
            t = re.sub(
                r'供应商名称[：:]\s*[xX]{3,4}\s*[（(]\s*单位[公盖]\s*[章]*\s*[）)][。.]?',
                f'供应商名称：{bidder}（盖章）。',
                t
            )
            t = re.sub(
                r'供应商名称[：:]\s*[xX]{3,4}\s*[（(]\s*盖\s*[章]*\s*[）)]',
                f'供应商名称：{bidder}（盖章）',
                t
            )
            t = re.sub(
                r'供应商名称[：:]\s*[xX]{3,4}\s*[（(]盖单位公章[）)]',
                f'供应商名称：{bidder}（盖章）',
                t
            )
            t = re.sub(
                r'(供\s*应\s*商名称[：:])\s*$',
                f'\\1{bidder}',
                t
            )

        # ==== 法定代表人/单位负责人或授权代表 ====
        if legal:
            t = re.sub(
                r'(法定代表人/单位负责人或授权代表)\s*[（\(]签字或加盖个人印章[）\)]\s*[：:]\s*[xX]{3,4}',
                f'\\1：{legal}',
                t
            )
            t = re.sub(
                r'(法定代表人/单位负责人或授权代表)\s*[（\(]签字或加盖个人印章[）\)]\s*[：:]\s*XXX+',
                f'\\1：{legal}',
                t
            )

        # ==== 授权书内容替换 ====
        if bidder and legal:
            t = re.sub(
                r'\s[Xx]{4}\s*[（\(]供应商名称[）\)]\s*[Xx]{4}\s*[（\(]法定代表人',
                f' {bidder}（供应商名称）{legal}（法定代表人',
                t
            )

        if authorized:
            t = re.sub(
                r'授权\s*[Xx]{4}\s*[（\(]被授权人',
                f'授权{authorized}（被授权人',
                t
            )

        # ==== 法定代表人/单位负责人签字 ====
        if legal:
            t = re.sub(
                r'(法定代表人/单位负责人（委托人）\s*签字或加盖个人印章[：:]\s*)[Xx]{3,4}',
                f'\\1{legal}',
                t
            )

        if authorized:
            t = re.sub(
                r'(授权代表（被授权人）\s*签字[：:]\s*)[Xx]{3,4}',
                f'\\1{authorized}',
                t
            )

        # ==== 通讯地址、邮编、电话、传真 ====
        if address:
            t = re.sub(r'通讯地址[：:]\s*[xX]{3,4}\s*$', f'通讯地址：{address}', t)
            t = re.sub(r'通讯地址[：:]\s*[xX]{3,4}(?:\s|$)', f'通讯地址：{address}', t)
        if zip_code:
            t = re.sub(r'邮政编码[：:]\s*[xX]{3,4}', f'邮政编码：{zip_code}', t)
        if phone:
            t = re.sub(r'联系电话[：:]\s*[xX]{3,4}[xX\-]*\s*$', f'联系电话：{phone}', t)
            t = re.sub(r'联系电话[：:]\s*[xX]{3,4}(?:\s|$)', f'联系电话：{phone}', t)
        if fax:
            t = re.sub(r'传\s*真[：:]\s*[xX]{3,4}', f'传    真：{fax}', t)

        # ==== 包号 ====
        t = t.replace('包        号：', f'包        号：{fill.get("package_no", "1")}')

        # ==== 采购项目编号：空值补填 ====
        if proj_id:
            t = re.sub(r'(采购项目编号[：:])\s*$', f'\\1{proj_id}', t)

        # ==== 最终兜底 ====
        t = re.sub(r'"[xX]{4,10}"', f'"{proj_name}"', t)
        t = re.sub(r'\u201c[xX]{4,10}\u201d', f'\u201c{proj_name}\u201d', t)

        return t


def generate_response_docx(bidding_docx_path: str, fill_data: dict) -> str:
    """便捷函数"""
    extractor = ResponseTemplateExtractor()
    return extractor.generate(bidding_docx_path, fill_data)
