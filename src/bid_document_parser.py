"""
招标文件解析器 v2 —— 精准提取项目信息、格式要求、开标日期
"""
import os
import re
import json
from typing import Dict, List, Optional


class BiddingDocumentParser:
    """解析招标文件，提取结构化信息"""

    def __init__(self):
        self.project_info: Dict[str, str] = {}
        self.requirements: Dict[str, any] = {}
        self.required_documents: List[str] = []
        self.format_requirements: Dict[str, str] = {}
        self.cover_requirements: Dict[str, str] = {}
        self.full_text: str = ""
        self._debug_extractions: Dict = {}

    def parse(self, file_path: str) -> Dict:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            self.full_text = self._parse_pdf(file_path)
        elif ext == '.docx':
            self.full_text = self._parse_docx(file_path)
        else:
            self.full_text = self._parse_txt(file_path)

        # ------ 提取流程 ------
        self._extract_project_info()       # 1. 项目名称、编号、招标人
        self._extract_dates()              # 2. 开标日期（核心）
        self._extract_agency()             # 2.5 采购代理机构
        self._extract_cover_requirements() # 3. 封面格式要求
        self._extract_document_order()     # 4. 投标文件组成及顺序
        self._extract_format_specs()       # 5. 装订/签署/密封要求
        self._extract_bidder_qualifications()  # 6. 资格要求
        self._extract_technical_requirements() # 7. 技术规格
        self._extract_evaluation_criteria()    # 8. 评标办法

        return self._build_result()

    # ==================== 文件解析 ====================

    def _parse_pdf(self, file_path: str) -> str:
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
            return text
        except:
            return self._parse_txt(file_path)

    def _parse_docx(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(c.text for c in row.cells)
            return text
        except:
            return self._parse_txt(file_path)

    def _parse_txt(self, file_path: str) -> str:
        for enc in ['utf-8', 'gbk', 'gb2312', 'gb18030']:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    return f.read()
            except:
                continue
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    # ==================== 1. 项目基本信息 ====================

    def _extract_project_info(self):
        text = self.full_text

        # ---- 项目名称：多策略匹配 ----
        name = None
        name_patterns = [
            # 精确匹配"项目名称："后的内容，跨行
            r'项目名称[：:\s]*[是为]?\s*(.+?)(?:\n|项目编号|招标编号|项目地点|采购预算)',
            # 采购项目名称
            r'采购项目名称[：:\s]*(.+?)(?:\n|项目编号|采购编号)',
            # "关于XXXX项目的招标公告" 模式
            r'关于[《]?\s*(.+?)\s*[》]?\s*(?:项目|工程|服务|采购).{0,10}(?:招标|采购|竞争性)',
            # 直接从表格行提取
            r'^\s*([^\s]{4,30}(?:项目|工程|采购|服务))\s+[A-Za-z0-9\-_]{3,}\s',
            # 文档开头的项目标题
            r'^[《]?\s*([^《》\n]{4,40}(?:项目|工程|采购|服务))\s*[》]?\s*(?:招标|采购|谈判)',
        ]
        for pat in name_patterns:
            m = re.search(pat, text, re.DOTALL | re.MULTILINE)
            if m:
                name = m.group(1).strip().rstrip('，。；;')
                # 清理多余字符
                name = re.sub(r'\s{2,}', ' ', name)
                if len(name) >= 2:
                    break

        # 兜底：取第一段中疑似项目名的内容
        if not name:
            lines = [l.strip() for l in text.split('\n')[:20] if l.strip()]
            for line in lines:
                if any(kw in line for kw in ['项目', '工程', '采购', '服务']):
                    # 提取引号或书名号中的内容
                    m = re.search(r'[《"\'](.+?)[》"\']', line)
                    if m and len(m.group(1)) >= 4:
                        name = m.group(1)
                        break
                    # 或直接取整行
                    if len(line) <= 50:
                        name = line
                        break

        if name:
            self.project_info['project_name'] = name
            self._debug_extractions['project_name_pattern'] = 'matched'

        # ---- 项目编号：放宽格式匹配 ----
        id_num = None
        id_patterns = [
            r'(?:项目编号|招标编号|采购编号|采购项目编号)[：:\s]*([A-Za-z0-9\-_]+)',
            r'编号[：:\s]*([A-Za-z0-9\-_]+)',
            r'(?:招标编号|项目编号).{0,3}([A-Za-z]+\d{4,}[A-Za-z0-9\-_]*)',
            r'(?:招标编号|项目编号).{0,3}([A-Z]{2,}[A-Za-z0-9\-_]{5,})',
        ]
        for pat in id_patterns:
            m = re.search(pat, text)
            if m:
                candidate = m.group(1).strip()
                if len(candidate) >= 3:
                    id_num = candidate
                    break

        if id_num:
            self.project_info['project_id'] = id_num
            self._debug_extractions['project_id_pattern'] = 'matched'

        # ---- 招标人 ----
        tenderer = None
        tenderer_patterns = [
            r'(?:招标人|采购人|招标单位|采购单位|采购人名称)[：:\s]*(.+?)(?:\n|地址|联系人|电话|$)',
            r'(?:招标人|采购人)[：:\s]*[名称]?[：:\s]*(.+?)(?:\n|地址|联系人|$)',
        ]
        for pat in tenderer_patterns:
            m = re.search(pat, text)
            if m:
                tenderer = m.group(1).strip().rstrip('，。')
                if tenderer:
                    break

        if tenderer:
            self.project_info['tenderee_name'] = tenderer

        # ---- 预算 ----
        budget = None
        budget_patterns = [
            r'(?:预算金额|最高限价|采购预算|项目预算|预算)[：:\s]*[人民币]?(\d[\d,.]*)\s*(?:万元?|元)',
            r'(?:预算金额|最高限价)[：:\s]*([\d,]+\.?\d*)\s*万',
        ]
        for pat in budget_patterns:
            m = re.search(pat, text)
            if m:
                budget = m.group(1).strip()
                break

        if budget:
            self.project_info['project_budget'] = budget

        # ---- 招标范围 ----
        scope = None
        scope_pat = r'(?:招标范围|采购内容|项目内容|建设内容|招标内容)[：:\s]*(.+?)(?:\n\n|\n(?:二|三|投标人|资格|技术|合同))'
        m = re.search(scope_pat, text, re.DOTALL)
        if m:
            scope = m.group(1).strip()[:300]
        if scope:
            self.project_info['project_scope'] = scope

    # ==================== 2. 日期提取（核心修复） ====================

    def _extract_dates(self):
        """提取投标截止时间/开标时间（用于封面日期）"""
        text = self.full_text

        # 开标时间（优先用于封面日期）
        opening = None
        opening_pats = [
            r'开标时间[：:\s]*(\d{4}\s*[年/\-]\s*\d{1,2}\s*[月/\-]\s*\d{1,2}\s*日?\s*\d{0,2}[：:]*\d{0,2})',
            r'开标时间[：:\s]*(.{8,30}?)(?:\n|投标|$)',
            r'开标日期[：:\s]*(\d{4}[年\-/]\d{1,2}[月\-/]\d{1,2}[日]?)',
        ]
        for pat in opening_pats:
            m = re.search(pat, text, re.DOTALL)
            if m:
                opening = m.group(1).strip().rstrip('。，；')
                break

        # 投标截止时间（备用）
        deadline = None
        deadline_pats = [
            r'(?:投标截止\s*(?:时间|日期|日?)|递交截止\s*(?:时间|日期|日?)|截止时间)[：:\s]*(\d{4}\s*[年/\-]\s*\d{1,2}\s*[月/\-]\s*\d{1,2}\s*日?\s*\d{0,2}[：:]*\d{0,2})',
            r'(?:投标截止|递交截止|截止时间)[：:\s]*(.{8,30}?)(?:\n|开标|$)',
            r'(\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日\s*\d{1,2}\s*[时:：]\s*\d{1,2}).{0,30}(?:截止|前|递交)',
        ]
        for pat in deadline_pats:
            m = re.search(pat, text, re.DOTALL)
            if m:
                deadline = m.group(1).strip().rstrip('。，；')
                break

        # 优先用开标日期，其次截止日期
        date_to_use = opening or deadline
        if date_to_use:
            # 统一格式化为中文日期
            self.project_info['bid_opening_date'] = self._normalize_date(date_to_use)
            self._debug_extractions['date_source'] = 'opening' if opening else 'deadline'

        # 另外单独存储两个日期
        if opening:
            self.requirements['bid_opening_time_raw'] = opening
        if deadline:
            self.requirements['submission_deadline_raw'] = deadline

        # 有效期
        validity = None
        m = re.search(r'投标有效期[：:\s]*(\d+)\s*[天日]', text)
        if m:
            validity = m.group(1)
        if validity:
            self.requirements['validity_days'] = validity

    def _normalize_date(self, raw: str) -> str:
        """将各种日期格式统一为 2024年12月15日"""
        raw = re.sub(r'\s+', '', raw)
        raw = re.sub(r'[时分秒：:]\d{1,2}[:：]\d{1,2}', '', raw)
        raw = re.sub(r'[时分秒：:]\d{1,2}', '', raw)
        
        m = re.match(r'(\d{4})[年\-/](\d{1,2})[月\-/](\d{1,2})[日]?', raw)
        if m:
            return f"{m.group(1)}年{m.group(2)}月{m.group(3)}日"
        return raw

    # ==================== 2.5 采购代理机构 ====================

    def _extract_agency(self):
        """提取采购代理机构名称"""
        text = self.full_text
        agency = None

        patterns = [
            r'(?:采购代理机构|招标代理机构|代理机构)[：:\s]*([^\n]{4,40}?)(?:\n|地址|联系人|电话|$)',
            r'([^\n]{4,30}(?:招标|代理|咨询|管理)(?:有限|股份)?公司)\s*受\s*[^\n]{2,20}\s*(?:委托|的委托)',
            r'(?:我司|我公司)\s*[：:在]?\s*([^\n]{4,30}(?:招标|代理|咨询|管理)(?:有限|股份)?公司)',
            r'代理机构[：:\s]*地址',
        ]
        for pat in patterns:
            m = re.search(pat, text)
            if m:
                agency = m.group(1).strip().rstrip('，。；')
                if agency and len(agency) >= 4:
                    break

        if agency:
            self.project_info['agency_name'] = agency
            self._debug_extractions['agency'] = 'found'

    # ==================== 3. 封面格式要求 ====================

    def _extract_cover_requirements(self):
        """从招标文件中提取封面/投标文件格式要求"""
        text = self.full_text

        cover_info = {}

        # 提取封面必须包含的内容
        cover_content_pats = [
            r'(?:封面|投标文件封面|投标文件外层信封).{0,20}应\s*(?:标注|标明|包含|注明|写明|载明)(.*?)(?:\n\n|\n(?:二|三|四|5\.|[0-9]+\.[0-9]))',
            r'(?:封面|投标文件封面).{0,10}要求[：:]*(.*?)(?:\n\n|\n(?!\s))',
        ]
        for pat in cover_content_pats:
            m = re.search(pat, text, re.DOTALL)
            if m:
                content = m.group(1).strip()
                # 拆分为独立要求项
                items = re.findall(r'[0-9]+[\.\)）]?\s*(.*?)(?=\n|$)', content)
                if items:
                    cover_info['required_items'] = items
                else:
                    cover_info['required_items'] = [content]
                break

        # 提取格式（如"正本/副本"标注位置）
        if '正本' in text or '副本' in text:
            # 查找份数要求
            m = re.search(r'(?:正本|副本|投标文件份数).{0,50}?(\d+)\s*份.*?(\d+)\s*份', text)
            if m:
                cover_info['copies'] = f"正本{m.group(1)}份，副本{m.group(2)}份"
            else:
                m = re.search(r'投标文件份数[：:]*\s*(.+?)(?:\n|$)', text)
                if m:
                    cover_info['copies'] = m.group(1).strip()

        self.cover_requirements = cover_info
        self._debug_extractions['cover_reqs_count'] = len(cover_info)

    # ==================== 4. 投标文件组成及顺序 ====================

    def _extract_document_order(self):
        """提取投标文件的组成章节及排列顺序"""
        text = self.full_text

        order_pats = [
            r'(?:投标文件\s*(?:组成|构成|应包括|应包含|内容|由以下|按以下顺序).{0,20})[：:]*\s*(.*?)(?:\n\n\s*(?:二|三|四|开标|评标|资格|合同|注|附件)|\n\s*\d+\.\s*[^0-9])',
            r'(?:投标人应提交|投标人须提交|需提供的材料|需递交的文件).{0,20}[：:]*\s*(.*?)(?:\n\n\s*(?:开标|评标|资格|合同|注|附件))',
        ]

        for pat in order_pats:
            m = re.search(pat, text, re.DOTALL)
            if m:
                content = m.group(1).strip()
                # 提取编号列表
                items = re.findall(
                    r'(?:^|\n)\s*(?:[（\(]?\s*(\d+)\s*[）\)\.、．]\s*)(.*?)(?=\n\s*(?:[（\(]?\s*\d+\s*[）\)\.、．]\s*)|\Z)',
                    content
                )
                if items:
                    # 按编号排序
                    self.required_documents = [
                        f"{num}. {desc.strip()}" for num, desc in sorted(items, key=lambda x: int(x[0]))
                    ]
                    self._debug_extractions['doc_order_count'] = len(self.required_documents)
                    return

                # Plan B: 用数字列表或破折号
                items = re.findall(r'(?:^|\n)\s*(?:[0-9]+[\.\)、．]|[-•◆■])\s*(.*?)(?=\n|$)', content)
                if items:
                    self.required_documents = [item.strip() for item in items if item.strip()]
                    return

        # 兜底默认顺序
        if not self.required_documents:
            self.required_documents = [
                "1. 投标函",
                "2. 开标一览表",
                "3. 法定代表人授权委托书",
                "4. 资格证明文件",
                "5. 技术方案",
                "6. 项目实施方案",
                "7. 类似项目业绩",
                "8. 其他材料"
            ]

    # ==================== 5. 格式规范 ====================

    def _extract_format_specs(self):
        """装订、签署、密封、字体等格式要求"""
        text = self.full_text

        specs = {
            'binding': None,    # 装订
            'signing': None,    # 签署/盖章
            'sealing': None,    # 密封
            'font_size': None,  # 字体字号
            'paper_size': None, # 纸张
            'copies': None,     # 份数
            'packaging': None,  # 包装
        }

        # 装订
        m = re.search(r'装订[：:]?\s*(.{0,60}?)(?:\n|。|$)', text)
        if m:
            specs['binding'] = m.group(1).strip()

        # 签署/盖章
        m = re.search(r'(?:签署|盖章|签章).{0,10}[：:]?\s*(.{0,80}?)(?:\n|。|$)', text)
        if m:
            specs['signing'] = m.group(1).strip()
        elif '加盖公章' in text or '逐页' in text:
            specs['signing'] = '投标文件须加盖单位公章，法定代表人或授权代表签字'

        # 密封
        m = re.search(r'密封[：:]?\s*(.{0,80}?)(?:\n|。|$)', text)
        if m:
            specs['sealing'] = m.group(1).strip()

        # 份数
        m = re.search(r'(?:份数|投标文件份数).{0,10}[：:]?\s*(.{0,60}?)(?:\n|。|$)', text)
        if m:
            specs['copies'] = m.group(1).strip()
        elif '正本' in text and '副本' in text:
            m = re.search(r'正本\s*(\d+)\s*份.*?副本\s*(\d+)\s*份', text)
            if m:
                specs['copies'] = f"正本{m.group(1)}份，副本{m.group(2)}份"

        self.format_requirements = {k: v for k, v in specs.items() if v}
        self._debug_extractions['format_specs_count'] = len(self.format_requirements)

    # ==================== 6. 资格要求 ====================

    def _extract_bidder_qualifications(self):
        text = self.full_text
        pats = [
            r'投标人资格要求[：:]?\s*(.*?)(?:\n\n\s*(?:二|三|四|技术|商务|评标|合同))',
            r'投标人资质条件[：:]?\s*(.*?)(?:\n\n\s*(?:二|三|技术|商务|评标))',
            r'资格条件[：:]?\s*(.*?)(?:\n\n\s*(?:技术|商务|评标))',
        ]
        for pat in pats:
            m = re.search(pat, text, re.DOTALL)
            if m:
                items = re.findall(r'(?:^|\n)\s*[0-9]+[\.\)）]\s*(.*?)(?=\n|$)', m.group(1))
                if items:
                    self.requirements['qualifications'] = items[:20]
                    return
                self.requirements['qualifications'] = [m.group(1).strip()[:500]]
                return

        self.requirements['qualifications'] = ["详见招标文件资格要求章节"]

    # ==================== 7. 技术要求 ====================

    def _extract_technical_requirements(self):
        text = self.full_text
        pats = [
            r'(?:技术\s*要求|技术规格|技术参数|服务要求|功能需求|建设内容)[：:]?\s*(.*?)(?:\n\n\s*(?:商务|评标|投标人|资格|合同|五|六))',
        ]
        for pat in pats:
            m = re.search(pat, text, re.DOTALL)
            if m:
                content = m.group(1).strip()
                items = re.findall(r'(?:^|\n)\s*[0-9]+[\.\)）]\s*(.*?)(?=\n|$)', content)
                if not items:
                    items = [content[:800]]
                self.requirements['technical'] = items[:15]
                return

        lines = re.findall(r'.{0,5}(?:须满足|须支持|须具备|应满足|应支持|应具备|技术要求).{10,100}', text)
        self.requirements['technical'] = lines[:10] if lines else ["详见招标文件技术要求章节"]

    # ==================== 8. 评标办法 ====================

    def _extract_evaluation_criteria(self):
        text = self.full_text
        pats = [
            r'(?:评标办法|评分标准|评审办法|评审标准|评标方法)[：:]?\s*(.*?)(?:\n\n\s*(?:合同|附件|投标人|注))',
        ]
        for pat in pats:
            m = re.search(pat, text, re.DOTALL)
            if m:
                items = re.findall(r'(?:^|\n)\s*[0-9]+[\.\)）]\s*(.*?)(?=\n|$)', m.group(1))
                if items:
                    self.requirements['evaluation_criteria'] = items[:20]
                    return
                self.requirements['evaluation_criteria'] = [m.group(1).strip()[:500]]
                return

        self.requirements['evaluation_criteria'] = ["详见招标文件评标办法章节"]

    # ==================== 汇总输出 ====================

    def _build_result(self) -> Dict:
        return {
            "project_info": self.project_info,
            "requirements": self.requirements,
            "required_documents": self.required_documents,
            "format_requirements": self.format_requirements,
            "cover_requirements": self.cover_requirements,
            "raw_text": self.full_text[:3000],
            "_debug": self._debug_extractions,
        }


def parse_bidding_document(file_path: str) -> Dict:
    parser = BiddingDocumentParser()
    return parser.parse(file_path)


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        result = parse_bidding_document(sys.argv[1])
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print("Usage: python bid_document_parser.py <bidding_document_path>")
